import fitz
import docx
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from rapidocr_onnxruntime import RapidOCR
import os
import glob
import numpy as np
from PIL import Image
import io
import re
import copy


PARAGRAPH_END_PATTERNS = re.compile(r'[。！？；]$')
STAMP_NOISE = re.compile(r'^[A-Z\-]{3,}$')
STAMP_OCR_NOISE = re.compile(r'^[A-Z0-9\-]{4,}$')


def set_chinese_font(run, font_name='宋体', font_size=12):
    run.font.size = Pt(font_size)
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def is_stamp_image(page, xref):
    rects = page.get_image_rects(xref)
    if not rects:
        return False
    rect = rects[0]
    page_w = page.rect.width
    page_h = page.rect.height
    img_w = rect.width
    img_h = rect.height
    area_ratio = (img_w * img_h) / (page_w * page_h)
    if area_ratio > 0.25:
        return False
    if img_w < 20 or img_h < 20:
        return False
    base_image = page.parent.extract_image(xref)
    ext = base_image.get('ext', '')
    if ext == 'jpeg':
        return False
    return True


def sample_stamp_color_from_render(page, rect):
    pix = page.get_pixmap(dpi=150, clip=rect)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    arr = np.array(img)
    non_white = arr[np.any(arr < 230, axis=2)]
    if len(non_white) == 0:
        return None
    r_vals = non_white[:, 0].astype(int)
    g_vals = non_white[:, 1].astype(int)
    b_vals = non_white[:, 2].astype(int)
    reddish = non_white[(r_vals > 120) & (g_vals < r_vals * 0.75) & (b_vals < r_vals * 0.75)]
    if len(reddish) > 0:
        r = int(reddish[:, 0].mean())
        g = int(reddish[:, 1].mean())
        b = int(reddish[:, 2].mean())
        return (r, g, b)
    return None


def extract_stamp_from_raw(pdf_doc, page, xref, temp_dir):
    rects = page.get_image_rects(xref)
    if not rects:
        return None
    rect = rects[0]

    stamp_color = sample_stamp_color_from_render(page, rect)
    if not stamp_color:
        stamp_color = (220, 50, 50)

    base_image = pdf_doc.extract_image(xref)
    img_bytes = base_image['image']
    stamp_img = Image.open(io.BytesIO(img_bytes))

    r, g, b = stamp_color

    if stamp_img.mode == '1':
        stamp_img = stamp_img.convert('L')
    if stamp_img.mode == 'P':
        stamp_img = stamp_img.convert('L')

    if stamp_img.mode == 'L':
        arr = np.array(stamp_img)
        dark_ratio = np.sum(arr < 128) / arr.size
        if dark_ratio > 0.5:
            stamp_mask = arr >= 128
        else:
            stamp_mask = arr < 128
        rgba = np.zeros((arr.shape[0], arr.shape[1], 4), dtype=np.uint8)
        rgba[stamp_mask, 0] = r
        rgba[stamp_mask, 1] = g
        rgba[stamp_mask, 2] = b
        rgba[stamp_mask, 3] = 200
        result_img = Image.fromarray(rgba, mode='RGBA')
    elif stamp_img.mode == 'RGBA':
        arr = np.array(stamp_img)
        gray = np.mean(arr[:, :, :3], axis=2)
        stamp_mask = gray < 128
        rgba = np.zeros_like(arr)
        rgba[:, :, 3] = 0
        rgba[stamp_mask, 0] = r
        rgba[stamp_mask, 1] = g
        rgba[stamp_mask, 2] = b
        rgba[stamp_mask, 3] = 200
        result_img = Image.fromarray(rgba, mode='RGBA')
    else:
        if stamp_img.mode != 'RGBA':
            stamp_img = stamp_img.convert('RGBA')
        arr = np.array(stamp_img)
        non_white = np.any(arr[:, :, :3] < 240, axis=2)
        arr[:, :, 3] = np.where(non_white, 200, 0)
        arr[non_white, 0] = r
        arr[non_white, 1] = g
        arr[non_white, 2] = b
        result_img = Image.fromarray(arr, mode='RGBA')

    page_w = page.rect.width
    page_h = page.rect.height
    display_width_cm = rect.width / page_w * 21.0
    display_height_cm = rect.height / page_h * 29.7
    rel_y = (rect.y0 + rect.y1) / 2 / page_h
    rel_x = (rect.x0 + rect.x1) / 2 / page_w

    temp_path = os.path.join(temp_dir, f"stamp_{xref}.png")
    result_img.save(temp_path, format='PNG')

    return {
        'path': temp_path,
        'rel_y': rel_y,
        'rel_x': rel_x,
        'display_width_cm': display_width_cm,
        'display_height_cm': display_height_cm,
        'rect': rect,
    }


def extract_stamp_images(pdf_doc, page, temp_dir):
    stamps = []
    images = page.get_images()
    for img_info in images:
        xref = img_info[0]
        if not is_stamp_image(page, xref):
            continue
        stamp = extract_stamp_from_raw(pdf_doc, page, xref, temp_dir)
        if stamp:
            stamps.append(stamp)
    stamps.sort(key=lambda s: s['rel_y'])
    return stamps


def add_floating_picture(paragraph, image_path, width_cm, pos_x_cm, pos_y_cm):
    run = paragraph.add_run()
    inline = run._element

    drawing_xml = f'''<w:drawing {nsdecls('wp', 'r', 'wp14', 'a', 'pic')}>
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251658240"
                 behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"
                 xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:posOffset>{int(pos_x_cm * 360000)}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>{int(pos_y_cm * 360000)}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{int(width_cm * 360000)}" cy="{int(width_cm * 360000)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="Picture 1"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="stamp.png"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{int(width_cm * 360000)}" cy="{int(width_cm * 360000)}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>'''

    return run


def ocr_lines_to_paragraphs(lines):
    if not lines:
        return []

    paragraphs = []
    current_para = [lines[0]]

    for i in range(1, len(lines)):
        prev = current_para[-1]
        curr = lines[i]

        prev_text = prev['text'].strip()
        curr_text = curr['text'].strip()

        prev_indent = prev['indent']
        curr_indent = curr['indent']
        prev_size = prev['font_size']
        curr_size = curr['font_size']

        prev_is_big = prev_size > 20
        curr_is_big = curr_size > 20

        is_new_para = False

        if prev_is_big and curr_is_big:
            is_new_para = False
        elif prev_is_big != curr_is_big:
            is_new_para = True
        elif curr['alignment'] == 'right' or prev['alignment'] == 'right':
            is_new_para = True
        elif curr_indent > 0.17 and prev_indent < 0.15:
            is_new_para = True
        elif curr_indent > 0.17 and abs(curr_indent - prev_indent) > 0.03:
            is_new_para = True
        elif PARAGRAPH_END_PATTERNS.search(prev_text) and curr_indent < 0.15:
            is_new_para = True

        if is_new_para:
            paragraphs.append(current_para)
            current_para = [curr]
        else:
            current_para.append(curr)

    if current_para:
        paragraphs.append(current_para)

    return paragraphs


def pdf_to_word_ocr(pdf_path, docx_path):
    ocr = RapidOCR()
    pdf_doc = fitz.open(pdf_path)
    word_doc = docx.Document()

    section = word_doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.9)
    section.right_margin = Cm(2.5)

    temp_dir = os.path.dirname(docx_path)

    all_paragraphs_data = []
    all_stamps = []

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]

        pix = page.get_pixmap(dpi=300)
        temp_img = os.path.join(temp_dir, f"temp_page_{page_num}.png")
        pix.save(temp_img)

        img = Image.open(temp_img)
        img_array = np.array(img)

        result, elapse = ocr(img_array)
        os.remove(temp_img)

        pt_to_cm = 2.54 / 72
        scale = pix.width / page.rect.width

        stamps = extract_stamp_images(pdf_doc, page, temp_dir)

        ocr_lines = []
        if result:
            for item in result:
                bbox = item[0]
                text = item[1]
                confidence = item[2]

                text = text.strip()
                if not text:
                    continue

                if STAMP_NOISE.match(text) and confidence < 0.85:
                    continue

                if STAMP_OCR_NOISE.match(text):
                    continue

                x_min = min(p[0] for p in bbox)
                y_min = min(p[1] for p in bbox)
                x_max = max(p[0] for p in bbox)
                y_max = max(p[1] for p in bbox)

                x_center = (x_min + x_max) / 2
                y_center = (y_min + y_max) / 2
                box_height = y_max - y_min
                box_width = x_max - x_min

                font_size = max(8, min(box_height * 72 / 300, 28))

                rel_y = y_center / pix.height
                indent = x_min / pix.width
                rel_x = x_center / pix.width
                right_margin = 1.0 - x_max / pix.width

                text_centered = abs(indent - right_margin) < 0.05 and indent > 0.1

                if indent > 0.5:
                    alignment = 'right'
                elif text_centered:
                    alignment = 'center'
                elif indent < 0.15:
                    alignment = 'left'
                elif indent > 0.17:
                    alignment = 'indent'
                else:
                    alignment = 'left'

                ocr_lines.append({
                    'text': text,
                    'indent': indent,
                    'rel_x': rel_x,
                    'rel_y': rel_y,
                    'font_size': font_size,
                    'alignment': alignment,
                    'page_num': page_num,
                    'left_indent_cm': indent * 21.0 - 2.9,
                    'pdf_y_cm': (y_min / scale) * pt_to_cm,
                    'pdf_y_end_cm': (y_max / scale) * pt_to_cm,
                })

        ocr_lines.sort(key=lambda x: x['rel_y'])

        para_groups = ocr_lines_to_paragraphs(ocr_lines)

        for para_lines in para_groups:
            avg_font_size = sum(l['font_size'] for l in para_lines) / len(para_lines)
            first_indent = para_lines[0]['indent']
            first_alignment = para_lines[0]['alignment']

            merged_text = ""
            for l in para_lines:
                if merged_text:
                    merged_text += l['text']
                else:
                    merged_text = l['text']

            all_paragraphs_data.append({
                'type': 'text',
                'text': merged_text.strip(),
                'font_size': avg_font_size,
                'indent': first_indent,
                'alignment': first_alignment,
                'page_num': page_num,
                'left_indent_cm': para_lines[0].get('left_indent_cm', 0),
                'pdf_y_cm': para_lines[0].get('pdf_y_cm', 0),
                'pdf_y_end_cm': para_lines[-1].get('pdf_y_end_cm', para_lines[-1].get('pdf_y_cm', 0)),
            })

        for stamp in stamps:
            all_stamps.append({
                'path': stamp['path'],
                'display_width_cm': stamp['display_width_cm'],
                'display_height_cm': stamp.get('display_height_cm', stamp['display_width_cm']),
                'rel_y': stamp['rel_y'],
                'rel_x': stamp['rel_x'],
                'page_num': page_num,
                'rect': stamp['rect'],
            })

    all_paragraphs_data.sort(key=lambda x: (x['page_num'], x.get('rel_y', 0)))

    stamp_paragraphs = {}

    prev_page = -1
    prev_pdf_y_end_cm = 0
    for idx, pdata in enumerate(all_paragraphs_data):
        if pdata['page_num'] != prev_page and prev_page != -1:
            word_doc.add_page_break()
            prev_pdf_y_end_cm = 0
        prev_page = pdata['page_num']

        text = pdata['text']
        font_size = pdata['font_size']
        alignment = pdata['alignment']
        indent = pdata['indent']
        pdf_y_cm = pdata.get('pdf_y_cm', 0)
        pdf_y_end_cm = pdata.get('pdf_y_end_cm', pdf_y_cm)

        if not text:
            continue

        is_date = bool(re.search(r'\d{4}年\d{1,2}月\d{1,2}日', text))
        is_signature = alignment == 'right' and not is_date and font_size < 20
        left_indent_cm = pdata.get('left_indent_cm', 0)
        is_contact = '联系人' in text or '联系电话' in text
        needs_vertical_pos = is_signature or is_date or is_contact

        if needs_vertical_pos and pdf_y_cm > 0 and prev_pdf_y_end_cm > 0:
            gap_cm = pdf_y_cm - prev_pdf_y_end_cm
            space_before_cm = max(0, gap_cm)
        else:
            space_before_cm = 0

        if font_size > 20 and not is_date:
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '方正小标宋简体', 22)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(16)
            p.paragraph_format.line_spacing = Pt(28)
        elif alignment == 'right' or is_date or is_signature:
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 16)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if left_indent_cm > 0:
                p.paragraph_format.left_indent = Cm(left_indent_cm)
            p.paragraph_format.space_before = Cm(space_before_cm) if space_before_cm > 0.1 else Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(28)
        elif alignment == 'center':
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(28)
        elif is_contact or (alignment == 'indent' and left_indent_cm > 1.5):
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 16)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(max(0, left_indent_cm))
            p.paragraph_format.space_before = Cm(space_before_cm) if space_before_cm > 0.1 else Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(28)
        else:
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 16)

            if alignment == 'indent' or indent > 0.17:
                p.paragraph_format.first_line_indent = Pt(32)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(28)

        prev_pdf_y_end_cm = pdf_y_end_cm

        for stamp in all_stamps:
            if stamp['page_num'] == pdata['page_num']:
                stamp_rel_y = stamp['rel_y']
                para_rel_y = pdata.get('rel_y', 0) if 'rel_y' in pdata else 0

    for stamp in all_stamps:
        target_para = None
        target_idx = -1

        for i, para in enumerate(word_doc.paragraphs):
            para_text = para.text.strip()
            has_left_indent = para.paragraph_format.left_indent and para.paragraph_format.left_indent > Cm(5)
            if has_left_indent and ('人民政府' in para_text or '局' in para_text or '委' in para_text or '办' in para_text or '公司' in para_text):
                target_para = para
                target_idx = i
                break

        if target_para is None:
            for i, para in enumerate(word_doc.paragraphs):
                if para.paragraph_format.left_indent and para.paragraph_format.left_indent > Cm(5):
                    target_para = para
                    target_idx = i
                    break

        if target_para is None:
            for i, para in enumerate(word_doc.paragraphs):
                para_text = para.text.strip()
                is_date = bool(re.search(r'\d{4}年\d{1,2}月\d{1,2}日', para_text))
                if is_date:
                    target_para = para
                    target_idx = i
                    break

        if target_para is None:
            target_para = word_doc.paragraphs[-1] if word_doc.paragraphs else word_doc.add_paragraph()

        run = target_para.add_run()
        width_cm = max(2.0, min(stamp['display_width_cm'], 5.0))
        height_cm = stamp.get('display_height_cm', width_cm)
        run.add_picture(stamp['path'], width=Cm(width_cm), height=Cm(height_cm))

        r_element = run._element
        drawing_elements = r_element.findall(qn('w:drawing'))

        if drawing_elements:
            drawing = drawing_elements[0]
            inline = drawing.find(qn('wp:inline'))

            if inline is not None:
                stamp_rect_pdf = stamp['rect']
                pdf_page = None
                for p_idx in range(len(pdf_doc)):
                    p = pdf_doc[p_idx]
                    if stamp['page_num'] == p_idx:
                        pdf_page = p
                        break

                if pdf_page:
                    pdf_pw = pdf_page.rect.width
                    pdf_ph = pdf_page.rect.height
                else:
                    pdf_pw = 595.0
                    pdf_ph = 841.0

                stamp_x0_cm = stamp_rect_pdf.x0 / pdf_pw * 21.0
                stamp_y0_cm = stamp_rect_pdf.y0 / pdf_ph * 29.7

                pos_x_cm = stamp_x0_cm
                pos_y_cm = stamp_y0_cm

                anchor_attribs = {
                    'distT': '0',
                    'distB': '0',
                    'distL': '0',
                    'distR': '0',
                    'simplePos': '0',
                    'relativeHeight': '251658240',
                    'behindDoc': '1',
                    'locked': '0',
                    'layoutInCell': '1',
                    'allowOverlap': '1',
                }
                anchor = inline.makeelement(qn('wp:anchor'), anchor_attribs)

                simple_pos = anchor.makeelement(qn('wp:simplePos'), {'x': '0', 'y': '0'})
                anchor.append(simple_pos)

                pos_h = anchor.makeelement(qn('wp:positionH'), {'relativeFrom': 'page'})
                pos_offset_h = pos_h.makeelement(qn('wp:posOffset'), {})
                pos_offset_h.text = str(int(pos_x_cm * 360000))
                pos_h.append(pos_offset_h)
                anchor.append(pos_h)

                pos_v = anchor.makeelement(qn('wp:positionV'), {'relativeFrom': 'page'})
                pos_offset_v = pos_v.makeelement(qn('wp:posOffset'), {})
                pos_offset_v.text = str(int(pos_y_cm * 360000))
                pos_v.append(pos_offset_v)
                anchor.append(pos_v)

                for child in list(inline):
                    tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag_local not in ('simplePos', 'positionH', 'positionV'):
                        anchor.append(child)

                drawing.remove(inline)
                drawing.append(anchor)

        os.remove(stamp['path'])

    word_doc.save(docx_path)
    pdf_doc.close()
    print(f"转换完成: {pdf_path} -> {docx_path}")


if __name__ == "__main__":
    pdf_files = glob.glob(os.path.join(r"d:\demo1", "*.pdf"))

    if not pdf_files:
        print("未找到PDF文件")
    else:
        for pdf_file in pdf_files:
            docx_file = os.path.splitext(pdf_file)[0] + ".docx"
            try:
                pdf_to_word_ocr(pdf_file, docx_file)
            except Exception as e:
                import traceback
                print(f"转换失败 {pdf_file}: {e}")
                traceback.print_exc()
